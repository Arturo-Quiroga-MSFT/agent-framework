"""Generate a Word document (architecture_components.docx) from the architecture_components.md file.

Minimal markdown-to-docx converter focused on the subset of syntax used in the
architecture document (headings, code fences, tables, paragraphs, inline code).

Requires: python-docx
    pip install python-docx
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError as e:  # pragma: no cover - handled at runtime
    raise SystemExit("python-docx not installed. Run: pip install python-docx") from e

ROOT = Path(__file__).resolve().parent
MD_FILE = ROOT / "architecture_components.md"
DOCX_FILE = ROOT / "architecture_components.docx"

CODE_FENCE = re.compile(r"^```(.*)$")
HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
TABLE_ROW = re.compile(r"^\|.*\|$")
INLINE_CODE = re.compile(r"`([^`]+)`")


def iter_lines(path: Path) -> Iterable[str]:
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            yield line.rstrip("\n")


def ensure_mono_style(doc: Document):
    styles = doc.styles
    if "CodeBlock" not in styles:
        style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = "Courier New"
        font.size = Pt(9)
        rpr = style._element.rPr
        rFonts = rpr.rFonts
        if rFonts is None:
            rFonts = rpr._add_rFonts()
        rFonts.set(qn('w:eastAsia'), 'Courier New')
    if "TableHeader" not in styles:
        style = styles.add_style("TableHeader", WD_STYLE_TYPE.PARAGRAPH)
        style.font.bold = True


def add_heading(doc: Document, text: str, level: int):
    doc.add_heading(text, level=level if level <= 4 else 4)


def add_code_block(doc: Document, code_lines: list[str], language: str | None):
    para = doc.add_paragraph(style="CodeBlock")
    # Join with newline preserving formatting
    para.add_run("\n".join(code_lines))


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(rows[0]))
    # Header
    hdr = table.add_row().cells
    for i, col in enumerate(rows[0]):
        hdr[i].text = col.strip()
    # Body
    for body_row in rows[1:]:
        cells = table.add_row().cells
        for i, col in enumerate(body_row):
            cells[i].text = col.strip()


def convert_md_to_docx():
    print(f"[debug] Working directory: {Path.cwd()}")
    print(f"[debug] Expected markdown path: {MD_FILE}")
    if not MD_FILE.exists():
        print("[error] Markdown source missing")
        raise SystemExit(f"Markdown file not found: {MD_FILE}")

    doc = Document()
    ensure_mono_style(doc)

    in_code = False
    code_lang: str | None = None
    code_buffer: list[str] = []
    table_buffer: list[list[str]] = []

    def flush_code():
        nonlocal code_buffer, code_lang
        if code_buffer:
            add_code_block(doc, code_buffer, code_lang)
            code_buffer = []
            code_lang = None

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

    for raw_line in iter_lines(MD_FILE):
        line = raw_line.rstrip()

        fence_match = CODE_FENCE.match(line)
        if fence_match:
            if in_code:
                # closing fence
                flush_code()
                in_code = False
            else:
                # starting fence
                in_code = True
                code_lang = fence_match.group(1).strip() or None
            continue

        if in_code:
            code_buffer.append(line)
            continue

        # Blank line => paragraph separation / flush table
        if not line.strip():
            flush_table()
            doc.add_paragraph("")
            continue

        # Headings
        h = HEADING.match(line)
        if h:
            flush_table()
            level = len(h.group(1))
            text = h.group(2).strip()
            add_heading(doc, text, level)
            continue

        # Tables
        if TABLE_ROW.match(line):
            # Split row on '|' ignoring first & last empty segments
            parts = [p.strip() for p in line.split("|")][1:-1]
            # Skip alignment row (---)
            if all(re.match(r"^:?-{3,}:?$", p) for p in parts):
                continue
            table_buffer.append(parts)
            continue

        # Regular paragraph (flush any table first)
        flush_table()
        para = doc.add_paragraph()
        last_end = 0
        # Inline code formatting
        for m in INLINE_CODE.finditer(line):
            if m.start() > last_end:
                para.add_run(line[last_end:m.start()])
            code_run = para.add_run(m.group(1))
            code_run.font.name = "Courier New"
            code_run.font.size = Pt(10)
            code_run.bold = False
            last_end = m.end()
        if last_end < len(line):
            para.add_run(line[last_end:])

    # Flush any trailing structures
    flush_code()
    flush_table()

    doc.save(DOCX_FILE)
    print(f"Generated {DOCX_FILE.relative_to(ROOT)} ({DOCX_FILE.stat().st_size} bytes)")


if __name__ == "__main__":  # pragma: no cover
    convert_md_to_docx()
